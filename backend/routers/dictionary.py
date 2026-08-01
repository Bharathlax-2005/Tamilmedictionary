import re
import csv
import json
import io
from fastapi import APIRouter, HTTPException, Depends, Query, UploadFile, File
from fastapi.responses import StreamingResponse, Response
from typing import Optional, List
from bson import ObjectId
from datetime import datetime, timezone
from ..database import get_db
from ..models import MedicalTermCreate, MedicalTermUpdate
from ..security import get_current_admin

router = APIRouter(prefix="/api/dictionary", tags=["Dictionary"])


def serialize_term(doc) -> dict:
    doc["id"] = str(doc["_id"])
    doc.pop("_id", None)
    return doc


@router.get("/search")
async def search_terms(
    q: str = Query("", min_length=0),
    page: int = Query(1, ge=1),
    limit: int = Query(20, ge=1, le=100),
    category: Optional[str] = None,
    featured: Optional[bool] = None,
):
    db = get_db()
    skip = (page - 1) * limit
    q_clean = q.strip()
    
    query = {}
    if category:
        query["category"] = category
    if featured is not None:
        query["is_featured"] = featured

    if q_clean:
        pattern = re.escape(q_clean)
        text_query = {
            "$or": [
                {"en_term": {"$regex": pattern, "$options": "i"}},
                {"ta_term": {"$regex": pattern, "$options": "i"}},
                {"definition": {"$regex": pattern, "$options": "i"}},
                {"ta_definition": {"$regex": pattern, "$options": "i"}},
                {"tags": {"$regex": pattern, "$options": "i"}},
                {"category": {"$regex": pattern, "$options": "i"}},
            ]
        }
        if query:
            query = {"$and": [query, text_query]}
        else:
            query = text_query

    total = await db.terms.count_documents(query)
    cursor = db.terms.find(query).skip(skip).limit(limit).sort("en_term", 1)
    terms = [serialize_term(doc) async for doc in cursor]
    return {
        "results": terms,
        "total": total,
        "page": page,
        "limit": limit,
        "pages": (total + limit - 1) // limit if total > 0 else 1,
    }


@router.get("/terms")
async def list_terms(
    page: int = Query(1, ge=1),
    limit: int = Query(20, ge=1, le=100),
    category: Optional[str] = None,
    featured: Optional[bool] = None,
):
    db = get_db()
    skip = (page - 1) * limit
    query = {}
    if category:
        query["category"] = category
    if featured is not None:
        query["is_featured"] = featured
    total = await db.terms.count_documents(query)
    cursor = db.terms.find(query).skip(skip).limit(limit).sort("en_term", 1)
    terms = [serialize_term(doc) async for doc in cursor]
    return {
        "results": terms,
        "total": total,
        "page": page,
        "limit": limit,
        "pages": (total + limit - 1) // limit if total > 0 else 1,
    }


@router.get("/categories")
async def get_categories():
    db = get_db()
    categories = await db.terms.distinct("category")
    clean_categories = sorted([c for c in categories if c])
    return {"categories": clean_categories}


@router.get("/export")
async def export_terms(
    format: str = Query("json", regex="^(json|csv)$"),
    _=Depends(get_current_admin),
):
    db = get_db()
    cursor = db.terms.find().sort("en_term", 1)
    terms = [serialize_term(doc) async for doc in cursor]

    if format == "json":
        json_data = json.dumps(terms, indent=2, default=str, ensure_ascii=False)
        return Response(
            content=json_data,
            media_type="application/json",
            headers={"Content-Disposition": 'attachment; filename="tamilmedictionary_terms.json"'},
        )
    else:
        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow(["en_term", "ta_term", "category", "definition", "ta_definition", "tags", "is_featured"])
        for t in terms:
            tags_str = ", ".join(t.get("tags", [])) if isinstance(t.get("tags"), list) else str(t.get("tags", ""))
            writer.writerow([
                t.get("en_term", ""),
                t.get("ta_term", ""),
                t.get("category", "General"),
                t.get("definition", ""),
                t.get("ta_definition", ""),
                tags_str,
                "true" if t.get("is_featured") else "false",
            ])
        output.seek(0)
        return Response(
            content=output.getvalue(),
            media_type="text/csv",
            headers={"Content-Disposition": 'attachment; filename="tamilmedictionary_terms.csv"'},
        )


HEADER_ALIASES = {
    "en_term": ["en_term", "english term", "english", "en", "term", "english_term"],
    "ta_term": ["ta_term", "tamil term", "tamil", "ta", "tamil translation", "tamil_term"],
    "category": ["category", "cat", "domain", "specialty"],
    "definition": ["definition", "english definition", "def", "english_definition", "description"],
    "ta_definition": ["ta_definition", "tamil definition", "ta_def", "tamil_definition"],
    "tags": ["tags", "tag", "keywords"],
    "is_featured": ["is_featured", "featured", "is featured", "feature"],
}


def normalize_header(header_name: str) -> Optional[str]:
    clean = str(header_name or "").strip().lower().replace("_", " ")
    for std_field, aliases in HEADER_ALIASES.items():
        for alias in aliases:
            if clean == alias or clean == alias.replace("_", " "):
                return std_field
    return None


def parse_csv_data(contents: bytes) -> List[dict]:
    decoded = contents.decode("utf-8-sig", errors="ignore")
    string_io = io.StringIO(decoded)
    sample = string_io.read(2048)
    string_io.seek(0)
    delimiter = ","
    try:
        dialect = csv.Sniffer().sniff(sample)
        delimiter = dialect.delimiter
    except Exception:
        pass

    reader = csv.reader(string_io, delimiter=delimiter)
    rows = list(reader)
    if not rows:
        return []

    header_row = rows[0]
    header_map = {}
    for col_idx, col_name in enumerate(header_row):
        norm = normalize_header(col_name)
        if norm:
            header_map[col_idx] = norm

    parsed_rows = []
    for r_idx, row in enumerate(rows[1:], start=2):
        if not any(row):
            continue
        row_dict = {}
        for c_idx, cell_value in enumerate(row):
            if c_idx in header_map:
                row_dict[header_map[c_idx]] = str(cell_value or "").strip()
        row_dict["_row_num"] = r_idx
        parsed_rows.append(row_dict)
    return parsed_rows


def parse_excel_data(contents: bytes) -> List[dict]:
    import openpyxl
    wb = openpyxl.load_workbook(filename=io.BytesIO(contents), data_only=True)
    sheet = wb.active
    rows = list(sheet.iter_rows(values_only=True))
    if not rows:
        return []

    header_row = None
    header_idx = 0
    for idx, r in enumerate(rows):
        if any(r):
            header_row = r
            header_idx = idx
            break

    if not header_row:
        return []

    header_map = {}
    for col_idx, cell in enumerate(header_row):
        norm = normalize_header(str(cell or ""))
        if norm:
            header_map[col_idx] = norm

    parsed_rows = []
    for r_idx, row in enumerate(rows[header_idx + 1:], start=header_idx + 2):
        if not any(row):
            continue
        row_dict = {}
        for c_idx, cell_value in enumerate(row):
            if c_idx in header_map:
                row_dict[header_map[c_idx]] = str(cell_value if cell_value is not None else "").strip()
        row_dict["_row_num"] = r_idx
        parsed_rows.append(row_dict)
    return parsed_rows


from ..unstructured_parser import UnstructuredTermExtractor


def parse_word_data(contents: bytes) -> List[dict]:
    # 1. Try structured table parsing first
    try:
        import docx
        doc = docx.Document(io.BytesIO(contents))
        if doc.tables:
            table = doc.tables[0]
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if len(rows) >= 2:
                header_row = rows[0]
                header_map = {}
                for col_idx, col_name in enumerate(header_row):
                    norm = normalize_header(col_name)
                    if norm:
                        header_map[col_idx] = norm
                if header_map:
                    parsed_rows = []
                    for r_idx, row in enumerate(rows[1:], start=2):
                        if not any(row):
                            continue
                        row_dict = {}
                        for c_idx, cell_value in enumerate(row):
                            if c_idx in header_map:
                                row_dict[header_map[c_idx]] = str(cell_value or "").strip()
                        row_dict["_row_num"] = r_idx
                        parsed_rows.append(row_dict)
                    if parsed_rows:
                        return parsed_rows
    except Exception:
        pass

    # 2. Fallback to Unstructured Text Extraction & Pattern Recognition
    raw_text = UnstructuredTermExtractor.extract_text_from_docx(contents)
    if not raw_text.strip():
        raise HTTPException(status_code=400, detail="No readable text found in Word document")

    parsed_unstructured = UnstructuredTermExtractor.parse_text_into_terms(raw_text)
    if not parsed_unstructured:
        raise HTTPException(status_code=400, detail="Could not extract structured medical terms from Word document text")
    return parsed_unstructured


def parse_pdf_data(contents: bytes) -> List[dict]:
    # 1. Try structured table parsing first
    try:
        import pdfplumber
        extracted_rows = []
        with pdfplumber.open(io.BytesIO(contents)) as pdf:
            for page in pdf.pages:
                tables = page.extract_tables()
                for t in tables:
                    for r in t:
                        if r and any(r):
                            clean_r = [str(cell or "").strip() for cell in r]
                            extracted_rows.append(clean_r)
        if len(extracted_rows) >= 2:
            header_row = extracted_rows[0]
            header_map = {}
            for col_idx, col_name in enumerate(header_row):
                norm = normalize_header(col_name)
                if norm:
                    header_map[col_idx] = norm
            if header_map:
                parsed_rows = []
                for r_idx, row in enumerate(extracted_rows[1:], start=2):
                    if not any(row):
                        continue
                    row_dict = {}
                    for c_idx, cell_value in enumerate(row):
                        if c_idx in header_map:
                            row_dict[header_map[c_idx]] = str(cell_value or "").strip()
                    row_dict["_row_num"] = r_idx
                    parsed_rows.append(row_dict)
                if parsed_rows:
                    return parsed_rows
    except Exception:
        pass

    # 2. Fallback to Unstructured PDF Text Extraction & Pattern Recognition
    raw_text = UnstructuredTermExtractor.extract_text_from_pdf(contents)
    if not raw_text.strip():
        raise HTTPException(status_code=400, detail="No readable text found in PDF document")

    parsed_unstructured = UnstructuredTermExtractor.parse_text_into_terms(raw_text)
    if not parsed_unstructured:
        raise HTTPException(status_code=400, detail="Could not extract structured medical terms from PDF document text")
    return parsed_unstructured


@router.post("/import")
async def import_terms(
    file: UploadFile = File(...),
    _=Depends(get_current_admin),
):
    db = get_db()
    contents = await file.read()
    filename = file.filename.lower()
    now = datetime.now(timezone.utc)

    if filename.endswith(".json"):
        try:
            raw_data = json.loads(contents.decode("utf-8"))
            if not isinstance(raw_data, list):
                raise HTTPException(status_code=400, detail="JSON file content must be a list of term objects")
            raw_rows = []
            for idx, item in enumerate(raw_data, start=1):
                item["_row_num"] = idx
                raw_rows.append(item)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to parse JSON file: {str(e)}")
    elif filename.endswith(".csv"):
        try:
            raw_rows = parse_csv_data(contents)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to parse CSV file: {str(e)}")
    elif filename.endswith(".xlsx") or filename.endswith(".xls"):
        try:
            raw_rows = parse_excel_data(contents)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to parse Excel file: {str(e)}")
    elif filename.endswith(".docx"):
        try:
            raw_rows = parse_word_data(contents)
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to parse Word document: {str(e)}")
    elif filename.endswith(".pdf"):
        try:
            raw_rows = parse_pdf_data(contents)
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to parse PDF document: {str(e)}")
    else:
        raise HTTPException(
            status_code=400,
            detail="Unsupported file format. Please upload a .csv, .xlsx, .xls, .docx, .pdf, or .json file."
        )


    if not raw_rows:
        raise HTTPException(status_code=400, detail="No readable rows or table data found in the uploaded file.")

    total_processed = len(raw_rows)
    inserted_count = 0
    updated_count = 0
    failed_count = 0
    errors = []

    for item in raw_rows:
        row_num = item.get("_row_num", 0)
        en = str(item.get("en_term", "")).strip()
        ta = str(item.get("ta_term", "")).strip()

        if not en:
            failed_count += 1
            errors.append({
                "row": row_num,
                "en_term": "",
                "reason": "Missing required English term (en_term)"
            })
            continue

        if not ta:
            ta = en

        category = str(item.get("category", "General")).strip() or "General"
        definition = str(item.get("definition", "")).strip() or None
        ta_definition = str(item.get("ta_definition", "")).strip() or None

        tags_raw = item.get("tags", [])
        if isinstance(tags_raw, str):
            tags = [t.strip() for t in tags_raw.split(",") if t.strip()]
        elif isinstance(tags_raw, list):
            tags = [str(t).strip() for t in tags_raw if str(t).strip()]
        else:
            tags = []

        is_feat_raw = item.get("is_featured", False)
        if isinstance(is_feat_raw, bool):
            is_featured = is_feat_raw
        else:
            is_featured = str(is_feat_raw).strip().lower() in ["true", "1", "yes"]

        term_doc = {
            "en_term": en,
            "ta_term": ta,
            "category": category,
            "definition": definition,
            "ta_definition": ta_definition,
            "tags": tags,
            "is_featured": is_featured,
            "updated_at": now,
        }

        existing = await db.terms.find_one({
            "en_term": {"$regex": f"^{re.escape(en)}$", "$options": "i"}
        })

        if existing:
            await db.terms.update_one(
                {"_id": existing["_id"]},
                {"$set": term_doc}
            )
            updated_count += 1
        else:
            term_doc["created_at"] = now
            await db.terms.insert_one(term_doc)
            inserted_count += 1

    return {
        "message": f"Processed {total_processed} rows: {inserted_count} inserted, {updated_count} updated, {failed_count} failed.",
        "total_processed": total_processed,
        "inserted": inserted_count,
        "updated": updated_count,
        "failed": failed_count,
        "errors": errors,
    }



@router.get("/terms/{term_id}")
async def get_term(term_id: str):
    db = get_db()
    if not ObjectId.is_valid(term_id):
        raise HTTPException(status_code=400, detail="Invalid ID")
    doc = await db.terms.find_one({"_id": ObjectId(term_id)})
    if not doc:
        raise HTTPException(status_code=404, detail="Term not found")
    return serialize_term(doc)


@router.post("/terms", status_code=201)
async def create_term(
    term: MedicalTermCreate,
    _=Depends(get_current_admin),
):
    db = get_db()
    doc = term.model_dump()
    doc["created_at"] = datetime.now(timezone.utc)
    doc["updated_at"] = datetime.now(timezone.utc)
    result = await db.terms.insert_one(doc)
    doc["id"] = str(result.inserted_id)
    doc.pop("_id", None)
    return doc


@router.put("/terms/{term_id}")
async def update_term(
    term_id: str,
    term: MedicalTermUpdate,
    _=Depends(get_current_admin),
):
    db = get_db()
    if not ObjectId.is_valid(term_id):
        raise HTTPException(status_code=400, detail="Invalid ID")
    update_data = {k: v for k, v in term.model_dump().items() if v is not None}
    update_data["updated_at"] = datetime.now(timezone.utc)
    result = await db.terms.update_one(
        {"_id": ObjectId(term_id)}, {"$set": update_data}
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Term not found")
    doc = await db.terms.find_one({"_id": ObjectId(term_id)})
    return serialize_term(doc)


@router.delete("/terms/{term_id}", status_code=204)
async def delete_term(term_id: str, _=Depends(get_current_admin)):
    db = get_db()
    if not ObjectId.is_valid(term_id):
        raise HTTPException(status_code=400, detail="Invalid ID")
    result = await db.terms.delete_one({"_id": ObjectId(term_id)})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Term not found")

