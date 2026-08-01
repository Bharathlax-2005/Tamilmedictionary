import re
import io

TAMIL_REGEX = re.compile(r'[\u0b80-\u0bff]+')


class UnstructuredTermExtractor:
    """
    Intelligent NLP & Pattern-recognition engine for extracting medical terms
    from unstructured paragraphs, headings, bullet points, and key-value blocks
    in Word (.docx) and PDF (.pdf) documents.
    """

    @staticmethod
    def extract_text_from_docx(contents: bytes) -> str:
        import docx
        doc = docx.Document(io.BytesIO(contents))
        text_parts = []

        # 1. Paragraphs & Headings
        for p in doc.paragraphs:
            txt = p.text.strip()
            if txt:
                text_parts.append(txt)

        # 2. Tables (if any table cells exist)
        for table in doc.tables:
            for row in table.rows:
                row_txt = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_txt:
                    text_parts.append(row_txt)

        return "\n".join(text_parts)

    @staticmethod
    def extract_text_from_pdf(contents: bytes) -> str:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(io.BytesIO(contents)) as pdf:
            for page in pdf.pages:
                txt = page.extract_text()
                if txt:
                    text_parts.append(txt)

                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        if row and any(row):
                            clean_row = " | ".join([str(c or "").strip() for c in row if str(c or "").strip()])
                            if clean_row:
                                text_parts.append(clean_row)

        return "\n".join(text_parts)

    @classmethod
    def parse_text_into_terms(cls, raw_text: str) -> list[dict]:
        lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
        parsed_terms = []
        current_block = {}

        for line_num, line in enumerate(lines, start=1):
            # 1. Key-Value pair lines: "Term: Hypertension", "Tamil: உயர் இரத்த அழுத்தம்", "Category: Diseases"
            kv_match = re.match(
                r'^(en_term|term|english|english term|ta_term|tamil|tamil term|category|cat|definition|def|english definition|tamil definition|ta_def|tags|tag|is_featured|featured)\s*[:=\-]\s*(.*)$',
                line,
                re.IGNORECASE
            )
            if kv_match:
                raw_key = kv_match.group(1).lower().strip().replace(" ", "_")
                val = kv_match.group(2).strip()

                if raw_key in ["term", "english", "english_term"]:
                    key = "en_term"
                elif raw_key in ["tamil", "tamil_term"]:
                    key = "ta_term"
                elif raw_key in ["cat"]:
                    key = "category"
                elif raw_key in ["def", "english_definition"]:
                    key = "definition"
                elif raw_key in ["tamil_definition", "ta_def"]:
                    key = "ta_definition"
                elif raw_key in ["tag"]:
                    key = "tags"
                elif raw_key in ["featured"]:
                    key = "is_featured"
                else:
                    key = raw_key

                # If starting a new term block and previous block is complete, flush it
                if key == "en_term" and "en_term" in current_block and "ta_term" in current_block:
                    current_block["_row_num"] = line_num - 1
                    parsed_terms.append(dict(current_block))
                    current_block = {}

                current_block[key] = val
                continue

            if "en_term" in current_block and "ta_term" in current_block:
                current_block["_row_num"] = line_num - 1
                parsed_terms.append(dict(current_block))
                current_block = {}

            # 2. Parentheses / Bracket Pattern: "Anatomy (உடலமைப்பியல்) - Study of body structure"
            parentheses_match = re.search(
                r'([A-Za-z0-9\s\-]+)[\(\[\{]([\u0b80-\u0bff\s]+)[\)\]\}](?:\s*[\-\:\;\|\–\—]\s*(.*))?',
                line
            )
            if parentheses_match:
                en = parentheses_match.group(1).strip()
                ta = parentheses_match.group(2).strip()
                df = parentheses_match.group(3).strip() if parentheses_match.group(3) else None
                if en and ta:
                    parsed_terms.append({
                        "en_term": en,
                        "ta_term": ta,
                        "definition": df,
                        "_row_num": line_num
                    })
                    continue


            # 3. Delimiter Separated Pattern: "Hypertension - உயர் இரத்த அழுத்தம்: High blood pressure"
            delimiters_match = re.split(r'\s*[\-\:\|\–\—\t]\s*', line)
            if len(delimiters_match) >= 2:
                p1 = delimiters_match[0].strip()
                p2 = delimiters_match[1].strip()
                p3 = delimiters_match[2].strip() if len(delimiters_match) >= 3 else None

                p1_has_tamil = bool(TAMIL_REGEX.search(p1))
                p2_has_tamil = bool(TAMIL_REGEX.search(p2))

                en_term = None
                ta_term = None
                definition = None

                if not p1_has_tamil and p2_has_tamil:
                    en_term = p1
                    ta_term = p2
                    definition = p3
                elif p1_has_tamil and not p2_has_tamil:
                    ta_term = p1
                    en_term = p2
                    definition = p3
                elif TAMIL_REGEX.search(line):
                    parts = cls._split_english_and_tamil(line)
                    if parts.get("en_term") and parts.get("ta_term"):
                        en_term = parts["en_term"]
                        ta_term = parts["ta_term"]
                        definition = parts.get("definition")

                if en_term and ta_term:
                    parsed_terms.append({
                        "en_term": en_term,
                        "ta_term": ta_term,
                        "definition": definition,
                        "_row_num": line_num
                    })
                    continue

            # 4. Fallback Script Separation Pattern for raw sentences
            if TAMIL_REGEX.search(line):
                parts = cls._split_english_and_tamil(line)
                if parts.get("en_term") and parts.get("ta_term"):
                    parts["_row_num"] = line_num
                    parsed_terms.append(parts)

        if "en_term" in current_block and "ta_term" in current_block:
            parsed_terms.append(current_block)

        return parsed_terms

    @staticmethod
    def _split_english_and_tamil(text: str) -> dict:
        tamil_match = TAMIL_REGEX.search(text)
        if not tamil_match:
            return {}

        ta_term = tamil_match.group(0).strip()

        # Replace Tamil script to find surrounding English term & definition
        remaining = TAMIL_REGEX.sub("||", text)
        parts = [p.strip(" :-|,()[]{}") for p in remaining.split("||") if p.strip(" :-|,()[]{}")]

        en_term = parts[0] if len(parts) > 0 else ""
        definition = parts[1] if len(parts) > 1 else None

        if en_term and ta_term:
            return {
                "en_term": en_term,
                "ta_term": ta_term,
                "definition": definition
            }
        return {}
