import requests
import io
import csv
import openpyxl
import docx

BASE_URL = "http://127.0.0.1:8000"

def get_auth_headers():
    login_res = requests.post(
        f"{BASE_URL}/api/auth/login",
        data={"username": "admin", "password": "Admin@1234"}
    )
    data = login_res.json()
    token = data.get("access_token")
    return {"Authorization": f"Bearer {token}"}


def create_csv_bytes():
    out = io.StringIO()
    writer = csv.writer(out)
    writer.writerow(["en_term", "ta_term", "category", "definition", "ta_definition", "tags", "is_featured"])
    writer.writerow(["Asthma", "ஆஸ்துமா", "Respiratory", "Chronic airway inflammation", "மூச்சுத்திணறல் நோய்", "respiratory, lung", "true"])
    writer.writerow(["Bronchitis", "மூச்சுக்குழாய் அழற்சி", "Respiratory", "Inflammation of bronchial tubes", "நுரையீரல் குழாய் அழற்சி", "respiratory, infection", "false"])
    return out.getvalue().encode("utf-8")

def create_excel_bytes():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["English Term", "Tamil Term", "Category", "Definition", "Tamil Definition", "Tags", "Featured"])
    ws.append(["Anemia", "இரத்த சோகை", "Hematology", "Lack of red blood cells", "இரத்த சிவப்பணு குறைவு", "blood, iron", True])
    ws.append(["Leukemia", "இரத்த புற்றுநோய்", "Hematology", "Cancer of blood-forming tissues", "இரத்த வெள்ளை அணு புற்றுநோய்", "blood, cancer", False])
    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()

def create_word_bytes():
    doc = docx.Document()
    table = doc.add_table(rows=1, cols=7)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "en_term"
    hdr_cells[1].text = "ta_term"
    hdr_cells[2].text = "category"
    hdr_cells[3].text = "definition"
    hdr_cells[4].text = "ta_definition"
    hdr_cells[5].text = "tags"
    hdr_cells[6].text = "is_featured"

    row_cells = table.add_row().cells
    row_cells[0].text = "Gastritis"
    row_cells[1].text = "இரைப்பை அழற்சி"
    row_cells[2].text = "Gastroenterology"
    row_cells[3].text = "Inflammation of stomach lining"
    row_cells[4].text = "வயிறு வீக்கம்"
    row_cells[5].text = "stomach, digestion"
    row_cells[6].text = "true"

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

def run_tests():
    headers = get_auth_headers()

    print("--- 1. Testing CSV Upload ---")
    csv_bytes = create_csv_bytes()
    files = {'file': ('terms_test.csv', csv_bytes, 'text/csv')}
    r = requests.post(f"{BASE_URL}/api/dictionary/import", files=files, headers=headers)
    print("CSV Response:", r.status_code, r.json())
    assert r.status_code == 200

    print("\n--- 2. Testing Excel Upload ---")
    excel_bytes = create_excel_bytes()
    files = {'file': ('terms_test.xlsx', excel_bytes, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')}
    r = requests.post(f"{BASE_URL}/api/dictionary/import", files=files, headers=headers)
    print("Excel Response:", r.status_code, r.json())
    assert r.status_code == 200

    print("\n--- 3. Testing Word Table Upload ---")
    word_bytes = create_word_bytes()
    files = {'file': ('terms_test.docx', word_bytes, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
    r = requests.post(f"{BASE_URL}/api/dictionary/import", files=files, headers=headers)
    print("Word Response:", r.status_code, r.json())
    assert r.status_code == 200

    print("\n--- 4. Testing Word Document Without Table (Should Reject) ---")
    empty_doc = docx.Document()
    empty_doc.add_paragraph("Just plain text without a table.")
    out = io.BytesIO()
    empty_doc.save(out)
    files = {'file': ('no_table.docx', out.getvalue(), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
    r = requests.post(f"{BASE_URL}/api/dictionary/import", files=files, headers=headers)
    print("No-Table Word Response:", r.status_code, r.json())
    assert r.status_code == 400

    res = requests.get(f"{BASE_URL}/api/dictionary/search?q=Gastritis")
    print("MongoDB query for 'Gastritis': status =", res.status_code)
    results = res.json().get("results", [])
    print("Found records count:", len(results))
    assert len(results) > 0
    assert results[0]["en_term"] == "Gastritis"



    print("\nSUCCESS: All multi-format upload and MongoDB upsert tests passed!")

if __name__ == "__main__":
    run_tests()
