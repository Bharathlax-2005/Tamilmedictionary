import requests
import io
import docx

BASE_URL = "http://127.0.0.1:8000"

def get_auth_headers():
    login_res = requests.post(
        f"{BASE_URL}/api/auth/login",
        data={"username": "admin", "password": "Admin@1234"}
    )
    token = login_res.json().get("access_token")
    return {"Authorization": f"Bearer {token}"}

def create_unstructured_word_bytes():
    doc = docx.Document()
    doc.add_heading("Medical Glossary Document (Unstructured Text)", level=1)
    doc.add_paragraph("Here are the latest medical term translations for the database:")
    doc.add_paragraph("Cardiomyopathy - இதயத்தசையகம் : Disease of the heart muscle")
    doc.add_paragraph("Encephalopathy (மூளை நோய்) - Disease affecting the brain")
    doc.add_paragraph("Otitis - செவி அழற்சி : Inflammation of the ear")
    doc.add_paragraph("Term: Nephrolithiasis")
    doc.add_paragraph("Tamil: சிறுநீரகக் கல்")
    doc.add_paragraph("Category: Nephrology")
    doc.add_paragraph("Definition: Kidney stones formation")
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

def run_unstructured_tests():
    headers = get_auth_headers()

    print("--- 1. Uploading Unstructured Word Document ---")
    word_bytes = create_unstructured_word_bytes()
    files = {'file': ('unstructured_terms.docx', word_bytes, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
    r = requests.post(f"{BASE_URL}/api/dictionary/import", files=files, headers=headers)
    print("Unstructured Word Import Response:", r.status_code, r.json())
    assert r.status_code == 200
    res_data = r.json()
    assert res_data.get("inserted", 0) + res_data.get("updated", 0) >= 3

    print("\n--- 2. Verifying MongoDB Upsert for Unstructured Terms ---")
    res = requests.get(f"{BASE_URL}/api/dictionary/search?q=Cardiomyopathy")
    assert res.status_code == 200
    results = res.json().get("results", [])
    print("Found 'Cardiomyopathy' in MongoDB:", len(results) > 0)
    assert len(results) > 0
    assert results[0]["en_term"] == "Cardiomyopathy"

    res_nephro = requests.get(f"{BASE_URL}/api/dictionary/search?q=Nephrolithiasis")
    assert res_nephro.status_code == 200
    results_nephro = res_nephro.json().get("results", [])
    print("Found 'Nephrolithiasis' in MongoDB:", len(results_nephro) > 0)
    assert len(results_nephro) > 0
    assert results_nephro[0]["en_term"] == "Nephrolithiasis"
    assert results_nephro[0]["category"] == "Nephrology"

    print("\nSUCCESS: Unstructured text parsing, structuring model, and MongoDB upsert verified end-to-end!")

if __name__ == "__main__":
    run_unstructured_tests()
